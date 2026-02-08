"""Word (.docx) document parser."""

from typing import List


def parse_word(docx_path: str, extract_image: bool = False) -> List[dict]:
    if extract_image:
        raise ValueError('Currently, extracting images is not supported!')

    from docx import Document
    doc = Document(docx_path)

    content = []
    for para in doc.paragraphs:
        content.append({'text': para.text})
    for table in doc.tables:
        tbl = []
        for row in table.rows:
            tbl.append('|' + '|'.join([cell.text for cell in row.cells]) + '|')
        tbl = '\n'.join(tbl)
        content.append({'table': tbl})

    return [{'page_num': 1, 'content': content}]
